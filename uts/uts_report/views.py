from django.contrib.auth.decorators import login_required
from django.shortcuts import render, get_object_or_404

from uts_common.models import Ticket


@login_required
def generate_ticket_report(request, pk):
    ticket = get_object_or_404(Ticket.objects.all(), pk=pk)
    from docx import Document
    from docx.shared import Inches
    document = Document()
    document.add_heading('Document Title', 0)
    p = document.add_paragraph('A plain paragraph having some ')
    p.add_run('bold').bold = True
    p.add_run(' and some ')
    p.add_run('italic.').italic = True
    document.add_heading('Heading, level 1', level=1)
    document.add_paragraph('Intense quote', style='Intense Quote')
    document.add_paragraph(
        'first item in unordered list', style='List Bullet'
    )
    document.add_paragraph(
        'first item in ordered list', style='List Number'
    )
    #document.add_picture('monty-truth.png', width=Inches(1.25))
    records = (
        (3, '101', 'Spam'),
        (7, '422', 'Eggs'),
        (4, '631', 'Spam, spam, eggs, and spam')
    )
    table = document.add_table(rows=1, cols=3)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Qty'
    hdr_cells[1].text = 'Id'
    hdr_cells[2].text = 'Desc'
    for qty, id, desc in records:
        row_cells = table.add_row().cells
        row_cells[0].text = str(qty)
        row_cells[1].text = id
        row_cells[2].text = desc
    document.add_page_break()
    from io import BytesIO
    target_stream = BytesIO()
    document.save(target_stream)
    from django.http import HttpResponse
    target_stream.seek(0)
    response = HttpResponse(target_stream, content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
    response["Content-Disposition"] = "attachment; filename=report.docx"
    return response