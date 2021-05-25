from docx import Document
from io import BytesIO
from docx.shared import Pt, Inches
import docx
import docx.oxml
import docx.opc.constants
from uts_common.models import TicketActions, TicketStatus


# https://github.com/python-openxml/python-docx/issues/384
def add_hyperlink(paragraph, url, text, color="0000FF", underline=True):
    """
    A function that places a hyperlink within a paragraph object.

    :param paragraph: The paragraph we are adding the hyperlink to.
    :param url: A string containing the required url
    :param text: The text displayed for the url
    :return: The hyperlink object
    """
    # This gets access to the document.xml.rels file and gets a new relation id value
    part = paragraph.part
    r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)
    # Create the w:hyperlink tag and add needed values
    hyperlink = docx.oxml.shared.OxmlElement('w:hyperlink')
    hyperlink.set(docx.oxml.shared.qn('r:id'), r_id, )
    # Create a w:r element
    new_run = docx.oxml.shared.OxmlElement('w:r')
    # Create a new w:rPr element
    rPr = docx.oxml.shared.OxmlElement('w:rPr')
    # Add color if it is given
    if not color is None:
      c = docx.oxml.shared.OxmlElement('w:color')
      c.set(docx.oxml.shared.qn('w:val'), color)
      rPr.append(c)
    # Remove underlining if it is requested
    if not underline:
      u = docx.oxml.shared.OxmlElement('w:u')
      u.set(docx.oxml.shared.qn('w:val'), 'none')
      rPr.append(u)
    # Join all the xml elements together add add the required text to the w:r element
    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)
    return hyperlink


def add_ticket_to_document(ticket, document, base_url, add_title=True):
    if add_title:
        document.add_heading(ticket.name, level=1)
    #p.add_run(f"Creato il {ticket.ts_open.strftime('%d/%m/%Y')} alle ore {ticket.ts_open.strftime('%H:%M:%S')}").italic = True
    for event in ticket.events.all():
        p = document.add_paragraph()
        p_owner = p.add_run(f"{event.owner}")
        p_owner.bold = True
        p_owner.font.size = Pt(10)
        p.add_run(f" ha ")
        p.add_run(f"{TicketActions[event.status]} ticket").italic = True
        p.add_run(" alle ore ")
        p_date = p.add_run(f"{event.timestamp.strftime('%H:%M:%S')}")
        p_date.italic = True
        p_date.font.size = Pt(10)
        p.add_run(" del ")
        p_time = p.add_run(f"{event.timestamp.strftime('%d/%m/%Y')}")
        p_time.italic = True
        p_time.font.size = Pt(10)
        p.paragraph_format.space_after = Pt(3)
        if event.status in [TicketStatus.ANSWER, TicketStatus.NOTE, TicketStatus.INFO_NEEDED] and event.info is not None:
            p = document.add_paragraph()
            p.paragraph_format.left_indent = Pt(9)
            p.paragraph_format.space_after = Pt(6)
            p_info = p.add_run(f"{event.info}")
            p_info.font.name = "Arial"
            p_info.font.size = Pt(10)
            for attachment in event.attachments.all():
                p = document.add_paragraph(style="List Bullet")
                hyperlink = add_hyperlink(p, f"{base_url}{attachment.file}", attachment.name)
    document.add_page_break()


# Generates a docx in memory and returns a stream of bytes which is the docx generated
def document_from_ticket(ticket, base_url):
    return document_from_many_tickets([ticket], f"{ticket.name}", base_url, add_title=False)


# Generates a docx in memory from many tickets
def document_from_many_tickets(tickets, title, base_url, add_title=True):
    document = Document()
    document.add_heading(f"{title}", 0)
    memory_stream = BytesIO()
    for ticket in tickets:
        add_ticket_to_document(ticket, document, base_url, add_title=add_title)
    document.save(memory_stream)
    memory_stream.seek(0)  # We need to get back at the start of the stream
    return memory_stream

